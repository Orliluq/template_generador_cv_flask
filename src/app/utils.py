import os
from flask import current_app
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert

def generate_cv():
    doc = Document()

    # Define estilos
    def set_heading(text, size=10, bold=True, color=RGBColor(0, 0, 0)):
        heading = doc.add_paragraph()
        run = heading.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = color
        run.font.name = 'Avenir'
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        heading.paragraph_format.space_before = Pt(6)
        heading.paragraph_format.space_after = Pt(6)
        return heading

    def add_section_title(text, size=11, bold=True, color=RGBColor(170, 59, 149)):
        section_title = doc.add_paragraph()
        run = section_title.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = color
        run.font.name = 'Avenir'
        section_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        section_title.paragraph_format.space_before = Pt(6)
        section_title.paragraph_format.space_after = Pt(2)
        return section_title

    def add_paragraph(text, size=10, bold=False, color=RGBColor(0, 0, 0), alignment=WD_PARAGRAPH_ALIGNMENT.LEFT):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = color
        run.font.name = 'Avenir'
        paragraph.alignment = alignment
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = Pt(12)  
        return paragraph
    
    def add_bulleted_list(doc, items, font_size=10, color=RGBColor(0, 0, 0)):
        for item in items:
            paragraph = doc.add_paragraph(style='List Bullet') 
            run = paragraph.add_run(item)
            run.font.size = Pt(font_size)
            run.font.color.rgb = color
            run.font.name = 'Avenir'
        
            paragraph.paragraph_format.space_before = Pt(0)  
            paragraph.paragraph_format.space_after = Pt(6)  
            paragraph.paragraph_format.line_spacing = Pt(12)  
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY 

    def add_projects_section():
        add_section_title("\nProyectos")

    # Header
    set_heading("Jhon Doe", size=16, bold=True)
    set_heading("Cargo Profesional Actual", size=14, bold=True)
    add_paragraph(
        "Ciudad, País | Email: jhondoe@example.com | GitHub: https://github.com/xxxx | LinkedIn: https://www.linkedin.com/in/xxxx | Portfolio: https://xxxxxxxxxx|",
        alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
    )

    # Profile Section
    add_section_title("\nPerfil")
    add_paragraph(
        "Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. "
        "Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum. "
        "Nemo enim ipsam voluptatem quia voluptas sit aspernatur aut odit aut fugit, sed quia consequuntur magni dolores eos qui ratione voluptatem sequi nesciunt.",
        alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    )

    # Skills Section
    add_section_title("\nHabilidades")
    add_bulleted_list(doc, [
        "Lorem Ipsum: is simply dummy text of the printing and typesetting industry",
        "Lorem Ipsum: dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua",
        "Lorem Ipsum: Excepteur sint occaecat cupidatat non proident ",
        "Herramientas y Otros: sunt in culpa qui officia deserunt mollit anim id est laborum"
    ])

    # Professional Experience Section
    add_section_title("\nExperiencia Profesional")

    # Experience 1
    add_paragraph(
        "1. Company Uno\nCargo-1\nFecha: xx/xxxx – Presente | Ciudad, País",
        bold=True,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    )
    add_bulleted_list(doc, [
        "Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et vas dolore magna aliqua.",
        "Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sinhtvt occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.",
        "Sed ut perspiciatis unde omnis iste natus error sit voluptatem accusantium doloremque laudantium, totam rem aperiam, eaque ipsa quae ab illo inventore veritatis et quasi architecto beatae vitae dicta sunt explicabo.",
    ])

    # Experience 2
    add_paragraph(
        "2. Company Dos\nCargo-2\nFecha: xx/xxxx – xx/xxxx | Ciudad, País",
        bold=True,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    )
    add_bulleted_list(doc, [
        "Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore uymagna aliqua.",
        "Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatuutr. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.",
        "Sed ut perspiciatis unde omnis iste natus error sit voluptatem accusantium doloremque laudantium, totam rem aperiam, eaque ipsa quae ab illo inventore veritatis et quasi architecto beatae vitae dicta sunt explicabo.",
    ])

    # Projects Section
    add_projects_section()

    # Proyecto 1
    add_paragraph(
        "1. Proyecto 1",
        bold=True,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    )
    add_bulleted_list(doc, [
        "Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore maougna aliqua.",
        "Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur siuint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.",
        "Enlace al Repositorio: https://xxxxxx",
        "Link: https://xxxxx/"
    ])

    # Proyecto 2
    add_paragraph(
        "2. Proyecto 2",
        bold=True,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    )
    add_bulleted_list(doc, [
        "Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.",
        "Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.",
        "Enlace al Repositorio: https://xxxxxx",
        "Link: https://xxxxx/"
    ])

    # Proyecto 3
    add_paragraph(
        "3. Proyecto 3",
        bold=True,
        alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
    )
    add_bulleted_list(doc, [
        "Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.",
        "Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.",
        "Enlace al Repositorio: https://xxxxxxxxx",
        "Link: https://xxxxxxxx/"
    ])

    # Education Section
    add_section_title("\nEducación")

    education_items = [
        "Lorem ipsum dolor sit amet\nsed do eiusmod tempor - Lorem ipsum Fecha: xx/xxxx – Presente | Remoto",
        "Duis aute irure dolor in reprehenderit\nHevfr Fecha: xx/xxxx – xx/xxxx | Presencial",
        "Lorem ipsum dolor sit amet\nsed do eiusmod tempor - Lorem ipsum Fecha: xx/xxxx – xx/xxxx | Híbrido",
    ]

    for item in education_items:
        paragraph = doc.add_paragraph(style='List Bullet')
        run = paragraph.add_run(item)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Avenir'
    
        paragraph.paragraph_format.space_before = Pt(0)  
        paragraph.paragraph_format.space_after = Pt(6)  
        paragraph.paragraph_format.line_spacing = Pt(12)  
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT 

    # Certifications Section
    add_section_title("\nCertificaciones")
    add_bulleted_list(doc, [
        "Lorem ipsum dolor sit amet, consectetur adipiscing elit, Certificado1, xxxx",
        "Excepteur sint occaecat cupidatat non proident, Certificado2, xxxx"
    ])

    # Languages Section
    add_section_title("\nIdiomas")
    add_paragraph("Excepteur: Lorem ipsum (x1).")

    output_dir = os.path.join(os.getcwd(), 'src', 'static')
    os.makedirs(output_dir, exist_ok=True)  
    docx_path = os.path.join(output_dir, 'CV.docx')
    pdf_path = os.path.join(output_dir, 'CV.pdf')

    doc.save(docx_path)
    # Comentando la conversión a PDF temporalmente
    # convert(docx_path, pdf_path)

    return docx_path, pdf_path  

generate_cv()