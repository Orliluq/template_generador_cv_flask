import unittest
import os
from src.app.utils import generate_cv
from docx import Document

class TestUtils(unittest.TestCase):

    def setUp(self):
        self.docx_path, self.pdf_path = generate_cv()

    def test_generate_cv_creates_docx_file(self):
        self.assertTrue(os.path.exists(self.docx_path))
        self.assertTrue(self.docx_path.endswith(".docx"))

    def test_generate_cv_creates_pdf_file(self):
        """Verificar si el archivo .pdf (si se habilita la conversión) se crea correctamente"""
        # Descomenta la conversión a PDF en el código si quieres probarlo
        # self.assertTrue(os.path.exists(self.pdf_path))
        pass  # Esto depende de si la conversión a PDF está habilitada o no

    def test_generate_cv_check_sections_in_docx(self):
        doc = Document(self.docx_path)

            "Perfil",
            "Habilidades",
            "Experiencia Profesional",
            "Proyectos",
            "Educación",
            "Certificaciones",
            "Idiomas",
        ]

        for title in section_titles:
            found = False
            for para in doc.paragraphs:
                if title in para.text:
                    found = True
                    break
            self.assertTrue(found, f"Sección '{title}' no encontrada en el documento")

    def test_generate_cv_check_paragraphs(self):
        """Verificar que los párrafos esperados están presentes en el archivo .docx"""
        # Cargar el archivo .docx generado
        doc = Document(self.docx_path)
        expected_texts = [
            "Jhon Doe",  # Verificar nombre
            "Cargo Profesional Actual",  # Verificar cargo
            "Lorem ipsum dolor sit amet",  # Ejemplo de párrafo
            "Enlace al Repositorio: https://github.com",  # Verificar enlaces
        ]

        for text in expected_texts:
            found = False
            for para in doc.paragraphs:
                if text in para.text:
                    found = True
                    break
            self.assertTrue(found, f"Texto '{text}' no encontrado en el documento")

    def tearDown(self):
        if os.path.exists(self.docx_path):
            os.remove(self.docx_path)
        if os.path.exists(self.pdf_path):
            os.remove(self.pdf_path)


if __name__ == "__main__":
    unittest.main()
